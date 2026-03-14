"""
Word exporter for Hong Kong university admission data.

Exports programme data to .docx format using python-docx.
"""

import logging
from pathlib import Path
from typing import List

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from hk_admissions.models import ProgramInfo

logger = logging.getLogger(__name__)


class WordExporter:
    """Export admission data to Word (.docx) format."""

    def export(
        self,
        programs: List[ProgramInfo],
        output_dir: str,
        filename_prefix: str = "hk_master_admissions",
    ) -> str:
        """
        Export programs to a Word document.

        Args:
            programs: List of ProgramInfo objects.
            output_dir: Directory to save the file.
            filename_prefix: Filename prefix.

        Returns:
            Path to the created Word file.
        """
        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(10)

        # Title
        title = doc.add_heading(
            "Hong Kong Universities Master's Programme Admissions", level=0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("香港各大学硕士招生信息汇总")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)

        doc.add_paragraph()  # spacer

        # Summary section
        doc.add_heading("Summary / 概览", level=1)
        uni_counts = {}
        for p in programs:
            key = p.university_name
            if key not in uni_counts:
                uni_counts[key] = {"count": 0, "cn": p.university_name_cn, "abbr": p.university_abbreviation}
            uni_counts[key]["count"] += 1

        summary_table = doc.add_table(rows=1, cols=4)
        summary_table.style = "Light Grid Accent 1"
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = "University"
        hdr_cells[1].text = "Chinese Name"
        hdr_cells[2].text = "Abbreviation"
        hdr_cells[3].text = "Programs"

        for name, data in sorted(uni_counts.items()):
            row = summary_table.add_row().cells
            row[0].text = name
            row[1].text = data["cn"]
            row[2].text = data["abbr"]
            row[3].text = str(data["count"])

        # Total
        total_row = summary_table.add_row().cells
        total_row[0].text = "TOTAL"
        total_row[3].text = str(len(programs))
        for cell in total_row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        doc.add_page_break()

        # Per-university detailed sections
        universities = {}
        for p in programs:
            key = p.university_name
            if key not in universities:
                universities[key] = []
            universities[key].append(p)

        for uni_name, uni_programs in sorted(universities.items()):
            uni_cn = uni_programs[0].university_name_cn if uni_programs else ""
            abbr = uni_programs[0].university_abbreviation if uni_programs else ""

            heading_text = f"{uni_name} ({abbr})"
            if uni_cn:
                heading_text += f" - {uni_cn}"
            doc.add_heading(heading_text, level=1)

            for idx, program in enumerate(uni_programs, 1):
                # Program heading
                doc.add_heading(
                    f"{idx}. {program.program_name}", level=2
                )

                # Details table
                details = [
                    ("Degree Type", program.degree_type),
                    ("Faculty / School", program.faculty),
                    ("Tuition Fee", program.tuition_fee or "Please refer to official website"),
                    ("Application Open Date", program.application_open_date or "Please refer to official website"),
                    ("Application Deadline", program.application_deadline or "Please refer to official website"),
                    ("Deadline Remarks", program.application_deadline_remarks),
                    ("English Requirement", program.english_requirement or "Please refer to official website"),
                    ("Duration", program.duration or "Please refer to official website"),
                    ("Mode", program.mode),
                    ("Official URL", program.program_url),
                    ("Remarks", program.remarks),
                ]

                # Filter out empty optional fields
                details = [(k, v) for k, v in details if v]

                table = doc.add_table(rows=len(details), cols=2)
                table.style = "Light List Accent 1"

                for row_idx, (label, value) in enumerate(details):
                    cells = table.rows[row_idx].cells
                    cells[0].text = label

                    # Make URL clickable
                    if label == "Official URL" and value:
                        paragraph = cells[1].paragraphs[0]
                        run = paragraph.add_run(value)
                        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                        run.font.underline = True
                    else:
                        cells[1].text = value

                    # Bold the label
                    for paragraph in cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                doc.add_paragraph()  # spacer

            if uni_name != sorted(universities.keys())[-1]:
                doc.add_page_break()

        # Footer note
        doc.add_paragraph()
        disclaimer = doc.add_paragraph()
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = disclaimer.add_run(
            "Disclaimer: All information is collected from official university websites. "
            "Please verify with the respective university for the most up-to-date information."
        )
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
        run.font.italic = True

        # Save
        filepath = str(Path(output_dir) / f"{filename_prefix}.docx")
        doc.save(filepath)
        logger.info(f"Word file saved: {filepath}")
        return filepath
