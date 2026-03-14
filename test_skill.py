#!/usr/bin/env python3
"""
Test script for Hong Kong Universities Master's Admissions Collector.

This script tests:
  1. Mock data generation and model validation
  2. All 5 export formats (Excel, Word, PDF, HTML, Markdown)
  3. Collector integration (with mock and optionally live scraping)
  4. File output integrity verification

Usage:
    # Run all tests with mock data (no network required)
    python test_skill.py

    # Run with live scraping from a specific university
    python test_skill.py --live --university hku

    # Run with live scraping from all universities
    python test_skill.py --live
"""

import os
import sys
import argparse
import logging
import traceback
from pathlib import Path
from datetime import datetime

# -----------------------------------------------------------------
# Setup
# -----------------------------------------------------------------
PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))

from hk_admissions.models import ProgramInfo, HK_UNIVERSITIES, UniversityConfig
from hk_admissions.exporters import export_all, EXPORTER_MAP
from hk_admissions.collector import HKAdmissionsCollector

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("test_skill")

OUTPUT_DIR = str(PROJECT_ROOT / "test_output")
FILENAME_PREFIX = "test_hk_admissions"

# -----------------------------------------------------------------
# ANSI color helpers
# -----------------------------------------------------------------
GREEN = "\033[92m"
RED = "\033[91m"
YELLOW = "\033[93m"
CYAN = "\033[96m"
BOLD = "\033[1m"
RESET = "\033[0m"


def ok(msg: str):
    print(f"  {GREEN}✅ PASS{RESET}: {msg}")


def fail(msg: str):
    print(f"  {RED}❌ FAIL{RESET}: {msg}")


def info(msg: str):
    print(f"  {CYAN}ℹ️  INFO{RESET}: {msg}")


def warn(msg: str):
    print(f"  {YELLOW}⚠️  WARN{RESET}: {msg}")


def section(title: str):
    print(f"\n{BOLD}{'='*70}")
    print(f"  {title}")
    print(f"{'='*70}{RESET}\n")


# -----------------------------------------------------------------
# Mock data factory
# -----------------------------------------------------------------
def create_mock_programs() -> list:
    """
    Create a realistic set of mock ProgramInfo objects covering all 9
    Hong Kong universities with diverse fields populated.
    """
    mock_data = [
        # HKU - 3 programs
        dict(
            university_name="The University of Hong Kong",
            university_name_cn="香港大学",
            university_abbreviation="HKU",
            faculty="Faculty of Business and Economics",
            faculty_cn="商学院",
            program_name="Master of Science in Business Analytics",
            program_name_cn="商业分析理学硕士",
            degree_type="MSc",
            tuition_fee="HKD 330,000",
            tuition_currency="HKD",
            application_open_date="2025-09-01",
            application_deadline="2026-04-14",
            application_deadline_remarks="Early round: Jan 13, 2026; Main round: Apr 14, 2026",
            english_requirement="IELTS 6.5 (no sub-score below 5.5) or TOEFL iBT 80",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://admissions.hku.hk/tpg/programme/master-of-science-in-business-analytics",
            data_source="https://admissions.hku.hk/tpg/programme-list",
        ),
        dict(
            university_name="The University of Hong Kong",
            university_name_cn="香港大学",
            university_abbreviation="HKU",
            faculty="Faculty of Engineering",
            program_name="Master of Science in Computer Science",
            program_name_cn="计算机科学理学硕士",
            degree_type="MSc",
            tuition_fee="HKD 210,000",
            application_open_date="2025-09-01",
            application_deadline="2026-01-13",
            english_requirement="IELTS 6.0 (no sub-score below 5.5) or TOEFL iBT 80",
            duration="1 year full-time / 2 years part-time",
            mode="Full-time / Part-time",
            program_url="https://admissions.hku.hk/tpg/programme/master-of-science-in-computer-science",
            data_source="https://admissions.hku.hk/tpg/programme-list",
        ),
        dict(
            university_name="The University of Hong Kong",
            university_name_cn="香港大学",
            university_abbreviation="HKU",
            faculty="Faculty of Law",
            program_name="Master of Laws (LLM)",
            degree_type="LLM",
            tuition_fee="HKD 178,200",
            application_deadline="2026-02-28",
            english_requirement="IELTS 7.0 (no sub-score below 6.5) or TOEFL iBT 97",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://www.law.hku.hk/postgraduate/llm/",
            data_source="https://admissions.hku.hk/tpg/programme-list",
        ),

        # CUHK - 2 programs
        dict(
            university_name="The Chinese University of Hong Kong",
            university_name_cn="香港中文大学",
            university_abbreviation="CUHK",
            faculty="Faculty of Engineering",
            program_name="MSc in Information Engineering",
            degree_type="MSc",
            tuition_fee="HKD 210,000",
            application_deadline="2026-02-28",
            english_requirement="IELTS 6.5 or TOEFL iBT 79",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://www.ie.cuhk.edu.hk/programmes/msc-in-information-engineering/",
            data_source="https://admissions.gs.cuhk.edu.hk/admissions/programme-list",
        ),
        dict(
            university_name="The Chinese University of Hong Kong",
            university_name_cn="香港中文大学",
            university_abbreviation="CUHK",
            faculty="CUHK Business School",
            program_name="Master of Accountancy",
            degree_type="MAcc",
            tuition_fee="HKD 300,000",
            application_deadline="2026-03-31",
            application_deadline_remarks="Round 1: Nov 2025; Round 2: Jan 2026; Round 3: Mar 2026",
            english_requirement="IELTS 6.5 or TOEFL iBT 79",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://macc.cuhk.edu.hk",
            data_source="https://admissions.gs.cuhk.edu.hk/admissions/programme-list",
        ),

        # HKUST - 2 programs
        dict(
            university_name="The Hong Kong University of Science and Technology",
            university_name_cn="香港科技大学",
            university_abbreviation="HKUST",
            faculty="School of Engineering",
            program_name="MSc in Big Data Technology",
            degree_type="MSc",
            tuition_fee="HKD 250,000",
            application_deadline="2026-03-01",
            english_requirement="IELTS 6.5 (no sub-score below 5.5) or TOEFL iBT 80",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://pg.ust.hk/prospective-students/programs/msc-bdt",
            data_source="https://pg.ust.hk/prospective-students/admissions/program-list",
        ),
        dict(
            university_name="The Hong Kong University of Science and Technology",
            university_name_cn="香港科技大学",
            university_abbreviation="HKUST",
            faculty="School of Business and Management",
            program_name="MBA Program",
            degree_type="MBA",
            tuition_fee="HKD 600,000",
            application_deadline="2026-03-31",
            application_deadline_remarks="R1: Oct 2025; R2: Jan 2026; R3: Mar 2026",
            english_requirement="IELTS 6.5 or TOEFL iBT 80; GMAT/GRE required",
            duration="16 months full-time",
            mode="Full-time",
            program_url="https://mba.hkust.edu.hk",
            data_source="https://pg.ust.hk/prospective-students/admissions/program-list",
        ),

        # PolyU - 2 programs
        dict(
            university_name="The Hong Kong Polytechnic University",
            university_name_cn="香港理工大学",
            university_abbreviation="PolyU",
            faculty="Faculty of Engineering",
            program_name="MSc in Electrical Engineering",
            degree_type="MSc",
            tuition_fee="HKD 171,000",
            application_open_date="2025-09-25",
            application_deadline="2026-04-30",
            english_requirement="IELTS 6.0 or TOEFL iBT 80",
            duration="1.5 years full-time",
            mode="Full-time / Part-time",
            program_url="https://www.polyu.edu.hk/eee/study/taught-postgraduate-programmes/msc-ee/",
            data_source="https://www.polyu.edu.hk/study/pg/taught-postgraduate-programmes",
        ),
        dict(
            university_name="The Hong Kong Polytechnic University",
            university_name_cn="香港理工大学",
            university_abbreviation="PolyU",
            faculty="School of Design",
            program_name="Master of Design",
            degree_type="MDes",
            tuition_fee="HKD 168,600",
            application_deadline="2026-04-30",
            english_requirement="IELTS 6.0 or TOEFL iBT 80",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://www.sd.polyu.edu.hk/en/study/master-of-design",
            data_source="https://www.polyu.edu.hk/study/pg/taught-postgraduate-programmes",
        ),

        # CityU - 2 programs
        dict(
            university_name="City University of Hong Kong",
            university_name_cn="香港城市大学",
            university_abbreviation="CityU",
            faculty="College of Business",
            program_name="MSc in Finance",
            degree_type="MSc",
            tuition_fee="HKD 282,000",
            application_deadline="2026-03-31",
            english_requirement="IELTS 6.5 or TOEFL iBT 79 or CET-6 490",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://www.cb.cityu.edu.hk/ef/programmes/mscfin/",
            data_source="https://www.admo.cityu.edu.hk/tpg/programmes",
        ),
        dict(
            university_name="City University of Hong Kong",
            university_name_cn="香港城市大学",
            university_abbreviation="CityU",
            faculty="College of Science",
            program_name="MSc in Data Science",
            degree_type="MSc",
            tuition_fee="HKD 198,000",
            application_deadline="2026-04-30",
            english_requirement="IELTS 6.5 or TOEFL iBT 79",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://www.cityu.edu.hk/ma/MSDataSci/",
            data_source="https://www.admo.cityu.edu.hk/tpg/programmes",
        ),

        # HKBU
        dict(
            university_name="Hong Kong Baptist University",
            university_name_cn="香港浸会大学",
            university_abbreviation="HKBU",
            faculty="School of Communication",
            program_name="MA in International Journalism Studies",
            degree_type="MA",
            tuition_fee="HKD 140,000",
            application_deadline="2026-04-05",
            english_requirement="IELTS 6.5 or TOEFL iBT 79",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://gs.hkbu.edu.hk/programmes/ma-in-international-journalism-studies",
            data_source="https://gs.hkbu.edu.hk/admission/taught-postgraduate-programmes",
        ),

        # EdUHK
        dict(
            university_name="The Education University of Hong Kong",
            university_name_cn="香港教育大学",
            university_abbreviation="EdUHK",
            faculty="Faculty of Education and Human Development",
            program_name="Master of Education",
            degree_type="MEd",
            tuition_fee="HKD 126,000",
            application_deadline="2026-03-01",
            english_requirement="IELTS 6.0 or TOEFL iBT 80",
            duration="1 year full-time",
            mode="Full-time / Part-time",
            program_url="https://www.eduhk.hk/gradsch/index.php/postgraduate-programmes/master-of-education/",
            data_source="https://www.eduhk.hk/acadprog/postgrad.html",
        ),

        # Lingnan
        dict(
            university_name="Lingnan University",
            university_name_cn="岭南大学",
            university_abbreviation="LU",
            faculty="Faculty of Business",
            program_name="MSc in Marketing and International Business",
            degree_type="MSc",
            tuition_fee="HKD 160,000",
            application_deadline="2026-05-31",
            english_requirement="IELTS 6.5 or TOEFL iBT 79",
            duration="1 year full-time",
            mode="Full-time",
            program_url="https://www.ln.edu.hk/mkt/mscmib/",
            data_source="https://www.ln.edu.hk/admissions/postgraduate/taught-postgraduate-programmes",
        ),

        # HKMU
        dict(
            university_name="Hong Kong Metropolitan University",
            university_name_cn="香港都会大学",
            university_abbreviation="HKMU",
            faculty="School of Science and Technology",
            program_name="MSc in Environmental Science and Management",
            degree_type="MSc",
            tuition_fee="HKD 109,200",
            application_deadline="2026-06-30",
            english_requirement="IELTS 6.0 or TOEFL iBT 79",
            duration="1 year full-time / 2 years part-time",
            mode="Full-time / Part-time",
            program_url="https://admissions.hkmu.edu.hk/pg/programmes/msc-environmental-science/",
            data_source="https://admissions.hkmu.edu.hk/pg/programmes/",
        ),
    ]

    programs = []
    for data in mock_data:
        p = ProgramInfo(
            last_updated=datetime.now().strftime("%Y-%m-%d"),
            **data,
        )
        programs.append(p)
    return programs


# -----------------------------------------------------------------
# Test functions
# -----------------------------------------------------------------
passed = 0
failed_count = 0


def assert_check(condition: bool, pass_msg: str, fail_msg: str):
    """Helper to track pass/fail counts."""
    global passed, failed_count
    if condition:
        ok(pass_msg)
        passed += 1
    else:
        fail(fail_msg)
        failed_count += 1


# -------------------- Test 1: Model & Validation --------------------
def test_model_validation():
    section("Test 1: ProgramInfo Model & Validation")

    # Test basic creation
    p = ProgramInfo(
        university_name="Test University",
        program_name="Test Program",
        program_url="https://example.com",
    )
    missing = p.validate()
    assert_check(
        len(missing) == 0,
        "ProgramInfo with all required fields passes validation",
        f"ProgramInfo validation failed unexpectedly: missing={missing}",
    )

    # Test missing fields
    p2 = ProgramInfo()
    missing2 = p2.validate()
    assert_check(
        len(missing2) == 3,
        f"Empty ProgramInfo correctly reports {len(missing2)} missing required fields: {missing2}",
        f"Expected 3 missing fields, got {len(missing2)}: {missing2}",
    )

    # Test to_dict
    d = p.to_dict()
    assert_check(
        isinstance(d, dict) and d["university_name"] == "Test University",
        "ProgramInfo.to_dict() returns correct dictionary",
        "ProgramInfo.to_dict() failed",
    )

    # Test UniversityConfig
    assert_check(
        len(HK_UNIVERSITIES) == 9,
        f"HK_UNIVERSITIES contains {len(HK_UNIVERSITIES)} universities (expected 9)",
        f"HK_UNIVERSITIES contains {len(HK_UNIVERSITIES)} universities, expected 9",
    )

    # Verify all university keys
    expected_keys = {"hku", "cuhk", "hkust", "polyu", "cityu", "hkbu", "eduhk", "lingnan", "hkmu"}
    actual_keys = set(HK_UNIVERSITIES.keys())
    assert_check(
        actual_keys == expected_keys,
        f"All 9 university keys are present: {sorted(actual_keys)}",
        f"University keys mismatch. Expected: {expected_keys}, Got: {actual_keys}",
    )


# -------------------- Test 2: Mock Data --------------------
def test_mock_data():
    section("Test 2: Mock Data Generation")

    programs = create_mock_programs()
    assert_check(
        len(programs) == 15,
        f"Created {len(programs)} mock programs (expected 15)",
        f"Expected 15 mock programs, got {len(programs)}",
    )

    # Check all 9 universities are represented
    uni_names = set(p.university_abbreviation for p in programs)
    assert_check(
        len(uni_names) == 9,
        f"Mock data covers {len(uni_names)} universities: {sorted(uni_names)}",
        f"Expected 9 universities, got {len(uni_names)}: {sorted(uni_names)}",
    )

    # Verify all programs pass validation
    all_valid = all(len(p.validate()) == 0 for p in programs)
    assert_check(
        all_valid,
        "All mock programs pass validation (required fields present)",
        "Some mock programs failed validation",
    )

    # Check key fields are populated
    has_tuition = sum(1 for p in programs if p.tuition_fee)
    has_deadline = sum(1 for p in programs if p.application_deadline)
    has_english = sum(1 for p in programs if p.english_requirement)
    has_url = sum(1 for p in programs if p.program_url)

    info(f"Fields populated: tuition={has_tuition}, deadline={has_deadline}, "
         f"english_req={has_english}, url={has_url}")

    assert_check(
        has_tuition == len(programs),
        f"All {has_tuition} programs have tuition fee info",
        f"Only {has_tuition}/{len(programs)} programs have tuition fee info",
    )
    assert_check(
        has_url == len(programs),
        f"All {has_url} programs have official URLs",
        f"Only {has_url}/{len(programs)} programs have official URLs",
    )

    return programs


# -------------------- Test 3: Excel Export --------------------
def test_excel_export(programs):
    section("Test 3: Excel Export (.xlsx)")

    from hk_admissions.exporters.excel_exporter import ExcelExporter

    exporter = ExcelExporter()
    try:
        filepath = exporter.export(programs, OUTPUT_DIR, FILENAME_PREFIX)
        file_exists = os.path.isfile(filepath)
        assert_check(
            file_exists,
            f"Excel file created: {filepath}",
            f"Excel file not found at: {filepath}",
        )

        if file_exists:
            size = os.path.getsize(filepath)
            assert_check(
                size > 1000,
                f"Excel file size: {size:,} bytes (reasonable size)",
                f"Excel file seems too small: {size} bytes",
            )

            # Verify with openpyxl
            from openpyxl import load_workbook
            wb = load_workbook(filepath)
            sheet_names = wb.sheetnames
            info(f"Sheets: {sheet_names}")

            assert_check(
                "All Programs" in sheet_names,
                "'All Programs' sheet exists",
                "'All Programs' sheet missing",
            )
            assert_check(
                "Summary" in sheet_names,
                "'Summary' sheet exists",
                "'Summary' sheet missing",
            )

            ws = wb["All Programs"]
            data_rows = ws.max_row - 1  # minus header
            assert_check(
                data_rows == len(programs),
                f"'All Programs' sheet has {data_rows} data rows (expected {len(programs)})",
                f"'All Programs' sheet has {data_rows} rows, expected {len(programs)}",
            )

            # Check per-university sheets
            uni_sheets = [s for s in sheet_names if s not in ("All Programs", "Summary")]
            info(f"Per-university sheets: {uni_sheets}")
            assert_check(
                len(uni_sheets) >= 9,
                f"Found {len(uni_sheets)} per-university sheets",
                f"Expected at least 9 per-university sheets, got {len(uni_sheets)}",
            )

            wb.close()

    except Exception as e:
        fail(f"Excel export error: {e}")
        traceback.print_exc()


# -------------------- Test 4: Word Export --------------------
def test_word_export(programs):
    section("Test 4: Word Export (.docx)")

    from hk_admissions.exporters.word_exporter import WordExporter

    exporter = WordExporter()
    try:
        filepath = exporter.export(programs, OUTPUT_DIR, FILENAME_PREFIX)
        file_exists = os.path.isfile(filepath)
        assert_check(
            file_exists,
            f"Word file created: {filepath}",
            f"Word file not found at: {filepath}",
        )

        if file_exists:
            size = os.path.getsize(filepath)
            assert_check(
                size > 1000,
                f"Word file size: {size:,} bytes (reasonable size)",
                f"Word file seems too small: {size} bytes",
            )

            # Verify structure with python-docx
            from docx import Document
            doc = Document(filepath)

            paragraphs = len(doc.paragraphs)
            tables = len(doc.tables)
            info(f"Paragraphs: {paragraphs}, Tables: {tables}")

            assert_check(
                paragraphs > 10,
                f"Word document has {paragraphs} paragraphs (content-rich)",
                f"Word document has only {paragraphs} paragraphs (too few)",
            )
            assert_check(
                tables >= 2,
                f"Word document has {tables} tables (summary + detail tables)",
                f"Word document has only {tables} tables (too few)",
            )

    except Exception as e:
        fail(f"Word export error: {e}")
        traceback.print_exc()


# -------------------- Test 5: PDF Export --------------------
def test_pdf_export(programs):
    section("Test 5: PDF Export (.pdf)")

    from hk_admissions.exporters.pdf_exporter import PDFExporter

    exporter = PDFExporter()
    try:
        filepath = exporter.export(programs, OUTPUT_DIR, FILENAME_PREFIX)
        file_exists = os.path.isfile(filepath)
        assert_check(
            file_exists,
            f"PDF file created: {filepath}",
            f"PDF file not found at: {filepath}",
        )

        if file_exists:
            size = os.path.getsize(filepath)
            assert_check(
                size > 1000,
                f"PDF file size: {size:,} bytes (reasonable size)",
                f"PDF file seems too small: {size} bytes",
            )

            # Basic PDF header check
            with open(filepath, "rb") as f:
                header = f.read(5)
            assert_check(
                header == b"%PDF-",
                "PDF file has valid PDF header (%PDF-)",
                f"PDF file has invalid header: {header}",
            )

    except Exception as e:
        fail(f"PDF export error: {e}")
        traceback.print_exc()


# -------------------- Test 6: HTML Export --------------------
def test_html_export(programs):
    section("Test 6: HTML Export (.html)")

    from hk_admissions.exporters.html_exporter import HTMLExporter

    exporter = HTMLExporter()
    try:
        filepath = exporter.export(programs, OUTPUT_DIR, FILENAME_PREFIX)
        file_exists = os.path.isfile(filepath)
        assert_check(
            file_exists,
            f"HTML file created: {filepath}",
            f"HTML file not found at: {filepath}",
        )

        if file_exists:
            size = os.path.getsize(filepath)
            assert_check(
                size > 1000,
                f"HTML file size: {size:,} bytes (reasonable size)",
                f"HTML file seems too small: {size} bytes",
            )

            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()

            assert_check(
                "<!DOCTYPE html>" in content,
                "HTML file contains DOCTYPE declaration",
                "HTML file missing DOCTYPE",
            )
            assert_check(
                "Hong Kong Universities" in content,
                "HTML file contains title content",
                "HTML file missing title content",
            )

            # Check that all universities appear
            for abbr in ["HKU", "CUHK", "HKUST", "PolyU", "CityU", "HKBU", "EdUHK", "LU", "HKMU"]:
                if abbr in content:
                    pass  # ok
                else:
                    warn(f"University abbreviation '{abbr}' not found in HTML content")

            assert_check(
                content.count("program-row") >= len(programs),
                f"HTML contains {content.count('program-row')} program rows (>= {len(programs)})",
                f"HTML has fewer program rows than expected",
            )

            # Check interactive elements
            assert_check(
                "searchInput" in content,
                "HTML includes search functionality",
                "HTML missing search functionality",
            )

    except Exception as e:
        fail(f"HTML export error: {e}")
        traceback.print_exc()


# -------------------- Test 7: Markdown Export --------------------
def test_markdown_export(programs):
    section("Test 7: Markdown Export (.md)")

    from hk_admissions.exporters.markdown_exporter import MarkdownExporter

    exporter = MarkdownExporter()
    try:
        filepath = exporter.export(programs, OUTPUT_DIR, FILENAME_PREFIX)
        file_exists = os.path.isfile(filepath)
        assert_check(
            file_exists,
            f"Markdown file created: {filepath}",
            f"Markdown file not found at: {filepath}",
        )

        if file_exists:
            size = os.path.getsize(filepath)
            assert_check(
                size > 500,
                f"Markdown file size: {size:,} bytes (reasonable size)",
                f"Markdown file seems too small: {size} bytes",
            )

            with open(filepath, "r", encoding="utf-8") as f:
                content = f.read()

            assert_check(
                "# " in content,
                "Markdown file contains headings",
                "Markdown file missing headings",
            )
            assert_check(
                "## " in content and "Table of Contents" in content,
                "Markdown file contains Table of Contents",
                "Markdown file missing Table of Contents",
            )
            assert_check(
                "| University |" in content,
                "Markdown file contains summary table",
                "Markdown file missing summary table",
            )
            assert_check(
                "Disclaimer" in content,
                "Markdown file contains disclaimer section",
                "Markdown file missing disclaimer",
            )

    except Exception as e:
        fail(f"Markdown export error: {e}")
        traceback.print_exc()


# -------------------- Test 8: export_all Integration --------------------
def test_export_all(programs):
    section("Test 8: export_all() Integration (All Formats)")

    try:
        results = export_all(
            programs=programs,
            output_dir=OUTPUT_DIR,
            filename_prefix=f"{FILENAME_PREFIX}_all",
            formats=None,  # all formats
        )

        assert_check(
            len(results) == 5,
            f"export_all returned {len(results)} format results (expected 5)",
            f"export_all returned {len(results)} results, expected 5",
        )

        all_success = all(v is not None for v in results.values())
        assert_check(
            all_success,
            "All 5 formats exported successfully",
            f"Some formats failed: {[k for k, v in results.items() if v is None]}",
        )

        for fmt, filepath in results.items():
            if filepath and os.path.isfile(filepath):
                ok(f"  {fmt}: {filepath} ({os.path.getsize(filepath):,} bytes)")
            else:
                fail(f"  {fmt}: file missing or export failed")

    except Exception as e:
        fail(f"export_all error: {e}")
        traceback.print_exc()


# -------------------- Test 9: Collector Structure --------------------
def test_collector_structure():
    section("Test 9: Collector Structure & Configuration")

    collector = HKAdmissionsCollector(max_workers=2)

    supported = collector.supported_universities
    assert_check(
        len(supported) == 9,
        f"Collector supports {len(supported)} universities",
        f"Collector supports {len(supported)} universities, expected 9",
    )

    for key, name in sorted(supported.items()):
        info(f"  {key}: {name}")

    # Test unknown university error
    try:
        collector.collect_university("nonexistent")
        fail("Expected ValueError for unknown university, but none raised")
    except ValueError as e:
        ok(f"Correctly raises ValueError for unknown university: {e}")

    # Test summary with empty data
    summary = collector.get_summary()
    assert_check(
        summary["total_programs"] == 0,
        "Empty collector returns zero programs in summary",
        f"Expected 0 programs in empty summary, got {summary['total_programs']}",
    )


# -------------------- Test 10: Live Scraping (Optional) --------------------
def test_live_scraping(university_key=None):
    section("Test 10: Live Scraping (Network Required)")

    collector = HKAdmissionsCollector(max_workers=3)

    if university_key:
        info(f"Scraping specific university: {university_key}")
        try:
            programs = collector.collect_university(university_key)
        except ValueError as e:
            fail(f"Invalid university key: {e}")
            return []
    else:
        info("Scraping ALL universities (this may take a while)...")
        programs = collector.collect_all()

    assert_check(
        len(programs) > 0,
        f"Live scraping collected {len(programs)} programs",
        "Live scraping collected 0 programs (might be network issue)",
    )

    summary = collector.get_summary()
    info(f"Total programs: {summary['total_programs']}")
    if summary.get("universities"):
        for uni, count in sorted(summary["universities"].items()):
            info(f"  {uni}: {count} programmes")

    if programs:
        # Export live data to all formats
        info("Exporting live data to all formats...")
        results = export_all(
            programs=programs,
            output_dir=OUTPUT_DIR,
            filename_prefix="live_hk_admissions",
            formats=None,
        )

        for fmt, filepath in results.items():
            if filepath and os.path.isfile(filepath):
                ok(f"  Live {fmt}: {filepath} ({os.path.getsize(filepath):,} bytes)")
            else:
                fail(f"  Live {fmt}: export failed")

    return programs


# -----------------------------------------------------------------
# Main
# -----------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="Test HK Admissions Collector Skill")
    parser.add_argument(
        "--live", action="store_true",
        help="Enable live scraping tests (requires network)",
    )
    parser.add_argument(
        "--university", "-u", default=None,
        help="Specific university key for live test (e.g., 'hku')",
    )
    args = parser.parse_args()

    print(f"\n{BOLD}{'#'*70}")
    print(f"#  HK Admissions Collector - Comprehensive Test Suite")
    print(f"#  Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"#  Output dir: {OUTPUT_DIR}")
    print(f"{'#'*70}{RESET}\n")

    # Ensure output directory exists
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Run tests
    test_model_validation()
    programs = test_mock_data()
    test_excel_export(programs)
    test_word_export(programs)
    test_pdf_export(programs)
    test_html_export(programs)
    test_markdown_export(programs)
    test_export_all(programs)
    test_collector_structure()

    if args.live:
        test_live_scraping(args.university)

    # Final summary
    print(f"\n{BOLD}{'='*70}")
    print(f"  FINAL RESULTS")
    print(f"{'='*70}{RESET}")
    total = passed + failed_count
    print(f"\n  Total tests:  {total}")
    print(f"  {GREEN}Passed:       {passed}{RESET}")
    print(f"  {RED}Failed:       {failed_count}{RESET}")

    if failed_count == 0:
        print(f"\n  {GREEN}{BOLD}🎉 ALL TESTS PASSED!{RESET}\n")
    else:
        print(f"\n  {RED}{BOLD}⚠️  {failed_count} TEST(S) FAILED{RESET}\n")

    # List generated files
    print(f"\n{BOLD}📁 Generated files in {OUTPUT_DIR}:{RESET}")
    if os.path.isdir(OUTPUT_DIR):
        for f in sorted(os.listdir(OUTPUT_DIR)):
            fpath = os.path.join(OUTPUT_DIR, f)
            if os.path.isfile(fpath):
                size = os.path.getsize(fpath)
                print(f"   📄 {f} ({size:,} bytes)")

    print()
    return 1 if failed_count > 0 else 0


if __name__ == "__main__":
    sys.exit(main())
